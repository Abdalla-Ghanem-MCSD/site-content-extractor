"""
builders.py
-----------
Turn the list of extracted pages into downloadable files:
  - Word document (.docx) with cover, headings and bullets
  - Markdown (.md)
  - JSON (.json) – raw structured data
"""

import json
from datetime import datetime
from urllib.parse import urlparse

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


NAVY = RGBColor(0x1F, 0x38, 0x64)
BLUE = RGBColor(0x2E, 0x54, 0x96)
GOLD = RGBColor(0xC9, 0xA2, 0x4B)
GREY = RGBColor(0x77, 0x77, 0x77)


def _site_name(url: str) -> str:
    return urlparse(url).netloc.replace("www.", "")


# --------------------------------------------------------------------------- #
# DOCX
# --------------------------------------------------------------------------- #
def build_docx(pages, out_path, source_url):
    doc = Document()

    # base font
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # tune heading colors
    for name, color, size in [("Heading 1", NAVY, 17),
                              ("Heading 2", BLUE, 14),
                              ("Heading 3", GOLD, 12)]:
        try:
            st = doc.styles[name]
            st.font.color.rgb = color
            st.font.size = Pt(size)
            st.font.bold = True
        except KeyError:
            pass

    # ---- cover ----
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(_site_name(source_url))
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = NAVY

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = s.add_run("Full Website Content Extract")
    rs.font.size = Pt(15)
    rs.font.color.rgb = GOLD
    rs.bold = True

    for line in [source_url,
                 f"{len(pages)} pages",
                 "Extracted " + datetime.now().strftime("%d %b %Y, %H:%M")]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = p.add_run(line)
        rr.font.size = Pt(10)
        rr.font.color.rgb = GREY

    doc.add_page_break()

    # ---- pages ----
    for i, page in enumerate(pages, 1):
        doc.add_heading(f"{i}. {page['title']}", level=1)

        url_p = doc.add_paragraph()
        ur = url_p.add_run("URL: " + page["url"])
        ur.italic = True
        ur.font.size = Pt(8)
        ur.font.color.rgb = GREY

        if page.get("meta_description"):
            mp = doc.add_paragraph()
            ml = mp.add_run("Meta description: ")
            ml.bold = True
            mp.add_run(page["meta_description"])

        for block in page["blocks"]:
            if block["type"] == "heading":
                lvl = min(max(block["level"], 1) + 1, 4)   # shift under page H1
                doc.add_heading(block["text"], level=lvl)
            elif block["type"] == "list":
                doc.add_paragraph(block["text"], style="List Bullet")
            else:
                doc.add_paragraph(block["text"])

        if i < len(pages):
            doc.add_page_break()

    doc.save(out_path)
    return out_path


# --------------------------------------------------------------------------- #
# Markdown
# --------------------------------------------------------------------------- #
def build_markdown(pages, out_path, source_url):
    lines = [f"# {_site_name(source_url)} — Website Content Extract",
             "",
             f"- Source: {source_url}",
             f"- Pages: {len(pages)}",
             f"- Extracted: {datetime.now().strftime('%d %b %Y, %H:%M')}",
             "", "---", ""]

    for i, page in enumerate(pages, 1):
        lines.append(f"## {i}. {page['title']}")
        lines.append("")
        lines.append(f"`{page['url']}`")
        lines.append("")
        if page.get("meta_description"):
            lines.append(f"> {page['meta_description']}")
            lines.append("")
        for block in page["blocks"]:
            if block["type"] == "heading":
                hashes = "#" * min(block["level"] + 2, 6)
                lines.append(f"{hashes} {block['text']}")
            elif block["type"] == "list":
                lines.append(f"- {block['text']}")
            else:
                lines.append(block["text"])
            lines.append("")
        lines.append("---")
        lines.append("")

    with open(out_path, "w", encoding="utf-8") as fh:
        fh.write("\n".join(lines))
    return out_path


# --------------------------------------------------------------------------- #
# JSON
# --------------------------------------------------------------------------- #
def build_json(pages, out_path, source_url):
    payload = {
        "source": source_url,
        "extracted_at": datetime.now().isoformat(),
        "page_count": len(pages),
        "pages": pages,
    }
    with open(out_path, "w", encoding="utf-8") as fh:
        json.dump(payload, fh, ensure_ascii=False, indent=2)
    return out_path
